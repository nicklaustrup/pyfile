import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
import docx
import os

class DocxViewer:
    """A viewer for DOCX documents."""
    def __init__(self, parent):
        self.docx_frame = ttk.Frame(parent)
        self.docx_frame.pack(fill="both", expand=True)

        # Create a header frame for the close button
        self.header_frame = ttk.Frame(self.docx_frame)
        self.header_frame.pack(fill="x")
        
        # Create close button but don't show it initially
        self.close_button = ttk.Button(self.header_frame, text="X", command=self.close_docx)
        # Don't pack the close button initially - it will be shown when a document is opened
        
        # Create a text widget to display the document content
        self.text_widget = scrolledtext.ScrolledText(
            self.docx_frame, 
            wrap=tk.WORD, 
            font=('Calibri', 11),
            padx=10,
            pady=10
        )
        self.text_widget.pack(fill="both", expand=True)
        
        # Make the text widget read-only
        self.text_widget.config(state=tk.DISABLED)
        
        # Initialize variables
        self.current_docx_path = None
        self.docx_document = None

    def open_docx(self, path):
        """Opens the DOCX document."""
        try:
            self.current_docx_path = path
            self.docx_document = docx.Document(path)
            
            # Show the close button
            self.close_button.pack(side="right")
            
            # Extract and display the document content
            self.display_document()
        except Exception as e:
            messagebox.showerror("Error", f"Error opening DOCX: {str(e)}")
    
    def display_document(self):
        """Extracts and displays the document content."""
        if not self.docx_document:
            return
        
        try:
            # Enable editing to update content
            self.text_widget.config(state=tk.NORMAL)
            
            # Clear existing content
            self.text_widget.delete(1.0, tk.END)
            
            # Extract and display paragraphs
            for para in self.docx_document.paragraphs:
                if para.text:
                    # Apply basic formatting
                    if para.style.name.startswith('Heading'):
                        self.text_widget.insert(tk.END, para.text + "\n", "heading")
                    else:
                        self.text_widget.insert(tk.END, para.text + "\n")
            
            # Extract and display tables
            for table in self.docx_document.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    self.text_widget.insert(tk.END, row_text + "\n")
                self.text_widget.insert(tk.END, "\n")
            
            # Disable editing after updating content
            self.text_widget.config(state=tk.DISABLED)
            
            # Scroll to the beginning
            self.text_widget.see("1.0")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error displaying document: {str(e)}")
    
    def close_docx(self):
        """Closes the DOCX document and resets the viewer."""
        # Enable editing to clear content
        self.text_widget.config(state=tk.NORMAL)
        self.text_widget.delete(1.0, tk.END)
        self.text_widget.config(state=tk.DISABLED)
        
        self.current_docx_path = None
        self.docx_document = None
        
        # Hide the close button
        self.close_button.pack_forget()
    
    def on_closing(self):
        """Handles the window close event."""
        if self.current_docx_path:
            response = messagebox.askyesnocancel("Close Document", "Do you want to close the current document?")
            if response:  # Yes
                self.close_docx()
                self.docx_frame.master.destroy()
            elif response is None:  # Cancel
                return
            else:  # No
                self.docx_frame.master.destroy()
        else:
            self.docx_frame.master.destroy() 